import streamlit as st
from docx import Document
import requests
import os

def translate_text(text, target_language):
    api_key = st.secrets["DASHSCOPE_API_KEY"]
    url = "https://dashscope-intl.aliyuncs.com/compatible-mode/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    prompt = f"Translate the following text to {target_language}: {text}"
    data = {
        "model": "qwen-turbo",
        "messages": [
            {
                "role": "system",
                "content": "You are a helpful assistant."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    }
    response = requests.post(url, headers=headers, json=data)
    if response.status_code == 200:
        translated_text = response.json()["choices"][0]["message"]["content"]
        return translated_text
    else:
        st.error(f"Error in translation: {response.text}")
        return None

def process_document(doc_path):
    try:
        doc = Document(doc_path)
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append(para)
        return paragraphs, doc
    except Exception as e:
        st.error(f"Error processing the document: {e}")
        return None, None

def save_translated_document(doc, corrected_paragraphs, output_path):
    try:
        for i, para in enumerate(doc.paragraphs):
            if i < len(corrected_paragraphs):
                para.text = corrected_paragraphs[i]
        doc.save(output_path)
    except Exception as e:
        st.error(f"Error saving the corrected document: {e}")

st.set_page_config(page_title="Text Correction Tool", page_icon="📝")
st.title("Spelling, Grammar, and Style Corrector 📝")

uploaded_file = st.file_uploader("Upload a Word document (.docx) for correction", type=["docx"])

if uploaded_file is not None:
    if uploaded_file.size == 0:
        st.error("The uploaded file is empty. Please upload a valid .docx file.")
        st.stop()

    temp_input_path = "temp_input.docx"
    try:
        with open(temp_input_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
    except Exception as e:
        st.error(f"Error saving the uploaded file: {e}")
        st.stop()

    paragraphs, doc = process_document(temp_input_path)
    
    if paragraphs is None or doc is None:
        st.error("The document could not be processed. Please ensure the file is a valid .docx document.")
        st.stop()

    original_language = st.text_input("Enter the primary language of the document (e.g., 'English'):", value="English")
    
    if st.button("Correct Document"):
        total_paragraphs = len(paragraphs)
        progress_bar = st.progress(0)
        corrected_paragraphs = []
        
        st.info("Performing spelling and grammar corrections...")
        for i, para in enumerate(paragraphs):
            corrected_text = translate_text(para.text, "English")
            if corrected_text:
                corrected_paragraphs.append(corrected_text)
                progress = (i + 1) / total_paragraphs
                progress_bar.progress(progress, text=f"Progress: {int(progress * 100)}%")
            else:
                st.warning(f"Failed to correct paragraph {i + 1}. Skipping...")
        
        if len(corrected_paragraphs) == total_paragraphs:
            st.success("First stage of correction completed.")
            
            st.info("Enhancing style and adjusting back to the original language...")
            final_corrected_paragraphs = []
            for i, para in enumerate(corrected_paragraphs):
                final_correction = translate_text(para, original_language)
                if final_correction:
                    final_corrected_paragraphs.append(final_correction)
                    progress = (i + 1) / total_paragraphs
                    progress_bar.progress(progress, text=f"Progress: {int(progress * 100)}%")
                else:
                    st.warning(f"Failed to finalize correction for paragraph {i + 1}. Skipping...")
            
            if len(final_corrected_paragraphs) == total_paragraphs:
                st.success("Second stage of correction completed.")
                
                temp_output_path = "corrected_output.docx"
                save_translated_document(doc, final_corrected_paragraphs, temp_output_path)
                
                try:
                    with open(temp_output_path, "rb") as f:
                        st.download_button(
                            label="Download Corrected Document",
                            data=f,
                            file_name="corrected_document.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"Error preparing the corrected document for download: {e}")
                
                try:
                    os.remove(temp_input_path)
                    os.remove(temp_output_path)
                except Exception as e:
                    st.warning(f"Could not clean up temporary files: {e}")
else:
    st.warning("Please upload a Word document (.docx) to proceed.")
